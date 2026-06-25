from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "SubLim3_Nexus_Owners_Manual.docx"
LOGO = ROOT / "core" / "public" / "assets" / "nexus-logo.png"

INK = "172033"
NAVY = "162B4D"
BLUE = "2E74B5"
CYAN = "25A7D9"
MUTED = "5B6576"
PALE = "EAF3F8"
LIGHT = "F3F6F8"
GOLD = "C48A1D"
RED = "9B1C1C"
WHITE = "FFFFFF"


def rgb(hex_value):
    return RGBColor.from_string(hex_value)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths_dxa, indent=120):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    marker = OxmlElement("w:tblHeader")
    marker.set(qn("w:val"), "true")
    tr_pr.append(marker)


def paragraph_border(paragraph, color=CYAN, size=18, space=8, side="left"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    edge = OxmlElement(f"w:{side}")
    edge.set(qn("w:val"), "single")
    edge.set(qn("w:sz"), str(size))
    edge.set(qn("w:space"), str(space))
    edge.set(qn("w:color"), color)
    p_bdr.append(edge)


def shade_paragraph(paragraph, fill=PALE):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def keep_with_next(paragraph, value=True):
    paragraph.paragraph_format.keep_with_next = value


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("SUBLIM3 NEXUS   |   ")
    run.font.name = "Aptos"
    run.font.size = Pt(8)
    run.font.color.rgb = rgb(MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)


def add_real_numbering(doc):
    numbering = doc.part.numbering_part.element
    existing_abs = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    existing_num = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_id = max(existing_abs or [0]) + 1
    num_id = max(existing_num or [0]) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start"); start.set(qn("w:val"), "1"); lvl.append(start)
    num_fmt = OxmlElement("w:numFmt"); num_fmt.set(qn("w:val"), "decimal"); lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText"); lvl_text.set(qn("w:val"), "%1."); lvl.append(lvl_text)
    suff = OxmlElement("w:suff"); suff.set(qn("w:val"), "tab"); lvl.append(suff)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs"); tab = OxmlElement("w:tab"); tab.set(qn("w:val"), "num"); tab.set(qn("w:pos"), "540"); tabs.append(tab); p_pr.append(tabs)
    ind = OxmlElement("w:ind"); ind.set(qn("w:left"), "540"); ind.set(qn("w:hanging"), "270"); p_pr.append(ind)
    spacing = OxmlElement("w:spacing"); spacing.set(qn("w:after"), "80"); spacing.set(qn("w:line"), "300"); spacing.set(qn("w:lineRule"), "auto"); p_pr.append(spacing)
    lvl.append(p_pr)
    abstract.append(lvl)
    numbering.append(abstract)
    num = OxmlElement("w:num"); num.set(qn("w:numId"), str(num_id))
    abs_id = OxmlElement("w:abstractNumId"); abs_id.set(qn("w:val"), str(abstract_id)); num.append(abs_id)
    numbering.append(num)
    return num_id


def apply_num(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl"); ilvl.set(qn("w:val"), "0")
    numid = OxmlElement("w:numId"); numid.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl); num_pr.append(numid); p_pr.append(num_pr)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.20

    title = styles["Title"]
    title.font.name = "Aptos Display"
    title.font.size = Pt(30)
    title.font.bold = True
    title.font.color.rgb = rgb(NAVY)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(8)

    subtitle = styles["Subtitle"]
    subtitle.font.name = "Aptos"
    subtitle.font.size = Pt(14)
    subtitle.font.color.rgb = rgb(MUTED)
    subtitle.paragraph_format.space_after = Pt(8)

    for style_name, size, color, before, after in (
        ("Heading 1", 19, NAVY, 18, 10),
        ("Heading 2", 14, BLUE, 14, 7),
        ("Heading 3", 11.5, NAVY, 10, 5),
    ):
        style = styles[style_name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "Step Label" not in styles:
        style = styles.add_style("Step Label", WD_STYLE_TYPE.PARAGRAPH)
    else:
        style = styles["Step Label"]
    style.font.name = "Aptos Display"; style.font.size = Pt(11); style.font.bold = True; style.font.color.rgb = rgb(NAVY)
    style.paragraph_format.space_before = Pt(6); style.paragraph_format.space_after = Pt(2); style.paragraph_format.keep_with_next = True

    if "Small Print" not in styles:
        style = styles.add_style("Small Print", WD_STYLE_TYPE.PARAGRAPH)
    else:
        style = styles["Small Print"]
    style.font.name = "Aptos"; style.font.size = Pt(8); style.font.color.rgb = rgb(MUTED)
    style.paragraph_format.space_after = Pt(3); style.paragraph_format.line_spacing = 1.05


def add_section_title(doc, number, title, kicker):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(f"{number}  /  {kicker.upper()}")
    r.bold = True; r.font.name = "Aptos"; r.font.size = Pt(9); r.font.color.rgb = rgb(CYAN)
    h = doc.add_paragraph(title, style="Heading 1")
    h.paragraph_format.space_before = Pt(0)
    return h


def add_callout(doc, label, text, kind="info"):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.15)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_together = True
    fill = PALE if kind == "info" else "FFF4DF" if kind == "caution" else "FBEAEA"
    color = CYAN if kind == "info" else GOLD if kind == "caution" else RED
    shade_paragraph(p, fill)
    paragraph_border(p, color=color, size=20, space=8)
    r = p.add_run(label.upper() + "  ")
    r.bold = True; r.font.color.rgb = rgb(color); r.font.size = Pt(9)
    r2 = p.add_run(text)
    r2.font.color.rgb = rgb(INK); r2.font.size = Pt(9.5)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.20
    return p


def add_numbered(doc, text, num_id):
    p = doc.add_paragraph(text)
    apply_num(p, num_id)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.20
    return p


def page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def credential_line(doc, label):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(label + "  ")
    r.bold = True; r.font.color.rgb = rgb(NAVY)
    p.add_run("________________________________________")
    return p


def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5); section.page_height = Inches(11)
    section.top_margin = Inches(0.78); section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.86); section.right_margin = Inches(0.86)
    section.header_distance = Inches(0.32); section.footer_distance = Inches(0.35)
    section.different_first_page_header_footer = True
    configure_styles(doc)
    num_id = add_real_numbering(doc)

    footer = section.footer.paragraphs[0]
    add_page_number(footer)

    # Cover - editorial handbook pattern.
    top = doc.add_paragraph()
    top.paragraph_format.space_before = Pt(18)
    top.paragraph_format.space_after = Pt(44)
    top.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if LOGO.exists():
        picture = top.add_run().add_picture(str(LOGO), width=Inches(1.05))
        picture._inline.docPr.set("descr", "SubLim3 Nexus logo")
        picture._inline.docPr.set("title", "SubLim3 Nexus")
    brand = doc.add_paragraph()
    brand.paragraph_format.space_after = Pt(2)
    r = brand.add_run("SUBLIM3")
    r.font.name = "Aptos Display"; r.font.size = Pt(14); r.bold = True; r.font.color.rgb = rgb(CYAN)
    title = doc.add_paragraph("Nexus", style="Title")
    title.paragraph_format.space_after = Pt(0)
    title.runs[0].font.size = Pt(42)
    subtitle = doc.add_paragraph("Owner's Manual", style="Subtitle")
    subtitle.runs[0].font.size = Pt(19)
    subtitle.runs[0].font.color.rgb = rgb(NAVY)
    rule = doc.add_paragraph()
    rule.paragraph_format.space_before = Pt(14); rule.paragraph_format.space_after = Pt(18)
    paragraph_border(rule, color=CYAN, size=22, space=1, side="top")
    lead = doc.add_paragraph("SETUP  •  OWNER / ADMIN  •  GAME MASTER  •  PLAYER")
    lead.runs[0].font.name = "Aptos"; lead.runs[0].font.size = Pt(10); lead.runs[0].bold = True; lead.runs[0].font.color.rgb = rgb(MUTED)
    lead.paragraph_format.space_after = Pt(30)
    copy = doc.add_paragraph("A local-first tabletop companion for campaigns, live encounters, player displays, audio, and RFID control.")
    copy.paragraph_format.right_indent = Inches(1.2)
    copy.runs[0].font.name = "Aptos Display"; copy.runs[0].font.size = Pt(16); copy.runs[0].font.color.rgb = rgb(INK)
    copy.paragraph_format.space_after = Pt(54)
    version = doc.add_paragraph("Product guide • Nexus Core 1.4.x")
    version.style = doc.styles["Small Print"]
    doc.core_properties.title = "SubLim3 Nexus Owner's Manual"
    doc.core_properties.subject = "Setup, Owner/Admin, Game Master, and Player instructions"
    doc.core_properties.author = "SubLim3"

    page_break(doc)
    add_section_title(doc, "01", "Before you begin", "Read this first")
    doc.add_paragraph("Nexus runs the table from a private local network. One Nexus Core stores campaign and character data, serves the Owner, GM, and Player screens, and coordinates media and RFID actions. Internet service is not required for normal local play.")
    doc.add_heading("Keep these items together", level=2)
    add_bullet(doc, "Nexus Core and the power supply provided for your unit.")
    add_bullet(doc, "The unit security card containing the Local Wi-Fi key, Owner recovery PIN, and guest GM PIN.")
    add_bullet(doc, "Any included RFID reader, cards, Player Controllers, and approved audio accessories.")
    add_callout(doc, "Credentials", "Treat the security card like a house key. Do not photograph or share the Owner recovery PIN. The guest GM PIN may be rotated at any time.", "caution")
    doc.add_heading("Placement and electrical safety", level=2)
    add_bullet(doc, "Use only the supplied or manufacturer-approved power adapter. Keep the cable where it cannot be pulled across the table.")
    add_bullet(doc, "Place Nexus on a stable, dry, ventilated surface. Do not cover ventilation openings or place the unit beside drinks, candles, or heat sources.")
    add_bullet(doc, "Do not open the enclosure. Disconnect power during electrical storms or before moving the unit.")
    add_bullet(doc, "Use Settings > Shut Down before removing power. Wait until the unit has fully stopped.")
    add_callout(doc, "Important", "Nexus is a tabletop accessory, not a toy. Small RFID cards and accessories may present a choking hazard. Keep them away from young children.", "warning")
    doc.add_heading("The three access levels", level=2)
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT; table.style = "Table Grid"
    headers = ["Owner / Admin", "Game Master", "Player"]
    descriptions = [
        "Full system access: campaigns, characters, clients, packs, media, RFID, settings, updates, and power.",
        "One campaign: scenes, encounters, initiative, health, conditions, and player invitations.",
        "One character: live resources, conditions, notes, scene, trackers, and turn status.",
    ]
    for i, cell in enumerate(table.rows[0].cells):
        cell.text = headers[i]; set_cell_shading(cell, NAVY); cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for run in cell.paragraphs[0].runs: run.font.color.rgb = rgb(WHITE); run.bold = True
    row = table.add_row()
    for i, cell in enumerate(row.cells): cell.text = descriptions[i]; cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    set_table_widths(table, [3120, 3120, 3120])
    set_repeat_table_header(table.rows[0])

    page_break(doc)
    add_section_title(doc, "02", "Quick start", "First power-on")
    add_callout(doc, "You will need", "A phone, tablet, or computer with Wi-Fi, plus the unit security card. If you plan to use a native app, install it before joining the Nexus Local Wi-Fi network.")
    steps = [
        ("Place and power Nexus", "Set Nexus on a ventilated surface and connect the supplied power adapter. Wait for the status light and allow the Core to finish starting."),
        ("Join Local Wi-Fi", "On your phone, tablet, or computer, join SubLim3-Nexus. Enter the Local Wi-Fi key printed on the unit security card. Your device may report that this network has no internet; stay connected."),
        ("Open the Owner Console", "In a browser, go to http://sublim3-nexus.local:3000. If the name does not resolve, use the IP address shown by your router or on the Settings network-status panel."),
        ("Connect as Owner", "Enter the Owner recovery PIN from the security card and choose Connect Owner Console. This browser remains paired for future visits; routine navigation should not ask for the PIN again."),
        ("Choose your network mode", "For direct offline play, keep Local Mode. To place Nexus on your existing network, open Settings, scan for Home Wi-Fi, select the network, and enter its password."),
        ("Create the table", "Install any optional Expansion Pack you need, create a campaign, add characters, then open GM Console. Share the campaign Player QR when everyone is ready."),
    ]
    for title_text, body in steps:
        p = add_numbered(doc, "", num_id)
        r = p.add_run(title_text + ". "); r.bold = True; r.font.color.rgb = rgb(NAVY)
        p.add_run(body)
    add_callout(doc, "Expected disconnect", "Changing Wi-Fi mode switches the Nexus radio and normally drops the current browser connection. Rejoin the selected network, then reopen sublim3-nexus.local:3000.", "caution")
    doc.add_heading("Record your unit information", level=2)
    credential_line(doc, "Serial / unit ID")
    credential_line(doc, "Local Wi-Fi key")
    credential_line(doc, "Owner recovery PIN")
    credential_line(doc, "Guest GM PIN")

    page_break(doc)
    add_section_title(doc, "03", "Owner / Admin setup", "Build your table")
    doc.add_paragraph("The Owner Console opens at the Nexus address. It is the only surface with full control of system configuration and stored table data.")
    doc.add_heading("1. Choose a game system", level=2)
    doc.add_paragraph("Open Expansion Packs. A new Nexus includes Custom RPG with eight quick-start hero presets. Install D&D 5e or another optional pack before creating a campaign that uses it. Packs are local and can be installed without changing existing campaigns.")
    add_callout(doc, "Pack removal", "A pack cannot be removed while a campaign uses it. Move or delete the dependent campaign first only if you are certain its data is no longer needed.", "caution")
    doc.add_heading("2. Create a campaign", level=2)
    add_numbered(doc, "Open Overview and scroll to Campaigns.", num_id)
    add_numbered(doc, "Enter a campaign name and select its game system.", num_id)
    add_numbered(doc, "Review the generated Campaign ID. IDs use lowercase letters, numbers, underscores, or hyphens and cannot contain spaces.", num_id)
    add_numbered(doc, "Choose Create campaign. The selected game system becomes the character template for this campaign.", num_id)
    doc.add_heading("3. Add player characters", level=2)
    add_numbered(doc, "In Characters, select the campaign.", num_id)
    add_numbered(doc, "Enter the character and player names. For Custom RPG, choose a ready-made hero or edit the generated fields.", num_id)
    add_numbered(doc, "Set template fields and resources, such as role, defense, health, mana, or other system-defined values.", num_id)
    add_numbered(doc, "Add starting conditions and public notes if needed. Public notes are visible on that character's Player screen.", num_id)
    add_numbered(doc, "Review the Character ID and choose Save character.", num_id)
    add_callout(doc, "Good practice", "Create every player character before displaying the Player QR. Players select from the campaign roster when they join.")

    page_break(doc)
    add_section_title(doc, "04", "Owner / Admin operation", "Access and system control")
    doc.add_heading("Owner access", level=2)
    doc.add_paragraph("The first successful Owner sign-in stores a local session in that browser. Owner sessions persist across restarts and normally expire after 90 days. Use Unpair this browser on a device you no longer control.")
    doc.add_heading("Guest GM access", level=2)
    add_numbered(doc, "In the Owner Console, locate Access & Pairing and note the guest GM PIN.", num_id)
    add_numbered(doc, "On the guest device, open /gm/, choose the campaign, enter the guest GM PIN, and pair the console.", num_id)
    add_numbered(doc, "When access is no longer needed, unpair the guest device or revoke it from the Owner Console.", num_id)
    add_callout(doc, "Rotate with care", "Rotating the GM PIN immediately revokes every current guest GM session. Owner browsers are not affected.", "caution")
    doc.add_heading("Paired clients and Player Controllers", level=2)
    doc.add_paragraph("The Access & Pairing panel lists connected clients, their role scope, and expiration date. Player Controllers also appear under Player Controllers with assigned campaign and character. Revoke or unpair a device that is lost, reassigned, or no longer trusted.")
    doc.add_heading("Live Session Overview", level=2)
    doc.add_paragraph("Select a campaign to inspect its published scene, mode, round, turn, and encounter status. The emergency reset clears the live session state; it does not delete the campaign or its characters.")
    doc.add_heading("Network and power", level=2)
    add_bullet(doc, "Local Mode broadcasts SubLim3-Nexus for direct, offline connection.")
    add_bullet(doc, "Home Wi-Fi joins an existing LAN. If connection fails, Nexus restores Local Mode.")
    add_bullet(doc, "Update downloads the latest configured release and restarts Nexus Core. Keep power connected until the result notice appears.")
    add_bullet(doc, "Reboot restarts the unit. Shut Down stops it safely; physical access is required to power it on again.")
    add_callout(doc, "Power safety", "Never disconnect power during an update. Use Shut Down before unplugging the Core.", "warning")

    page_break(doc)
    add_section_title(doc, "05", "Media and RFID", "Build atmosphere")
    doc.add_heading("Add and organize audio", level=2)
    add_numbered(doc, "Open Media Library. Create folders to organize music, ambience, and effects.", num_id)
    add_numbered(doc, "Choose Add audio to upload supported files from the current device, or mount a USB drive and choose Scan USB.", num_id)
    add_numbered(doc, "Review imported files, names, and folders before play.", num_id)
    doc.add_heading("Run table audio", level=2)
    doc.add_paragraph("Open Media to search Nexus storage, mounted USB media, or live radio. Select ambience, use the transport controls, adjust volume, or trigger one-shot effects. Live radio requires internet access; local files and built-in procedural content do not.")
    add_callout(doc, "Audio output", "The current browser provides the playback renderer. Keep the Media page open on the device connected to your table speakers while audio is playing.")
    doc.add_heading("Bind an RFID card", level=2)
    add_numbered(doc, "Open RFID Cards and scan the card on the connected reader. You may also enter its hexadecimal UID.", num_id)
    add_numbered(doc, "Choose Use last scan, give the card a recognizable name, and select an action.", num_id)
    add_numbered(doc, "For Play audio or effect, choose the library item. Function cards may Stop, Pause, Volume up, or Volume down.", num_id)
    add_numbered(doc, "Save the binding, then use Test entered UID or scan the physical card to verify it.", num_id)
    doc.add_heading("Choose card behavior", level=2)
    doc.add_paragraph("In Settings, choose Swipe to start playback or Place to play only while a card is present. Choose what a second scan does (toggle, restart, or ignore), set a repeat-scan delay, and decide whether function cards bypass that delay.")
    add_callout(doc, "Volume limit", "Set Maximum volume and Startup volume in Settings before the first session. Startup volume cannot be higher than the maximum.", "caution")

    page_break(doc)
    add_section_title(doc, "06", "Game Master directions", "Prepare and invite")
    doc.add_heading("Open the GM Console", level=2)
    doc.add_paragraph("An Owner browser opens GM Console automatically and may switch among campaigns. A guest GM opens /gm/, selects one campaign, and enters the guest GM PIN. A guest GM remains limited to the paired campaign.")
    doc.add_heading("Invite players", level=2)
    add_numbered(doc, "Choose the active campaign in GM Console.", num_id)
    add_numbered(doc, "Show the Player QR, or use Copy link / Share to send the local campaign link.", num_id)
    add_numbered(doc, "Ask each player to connect to the same Nexus network, scan the code, and choose their character.", num_id)
    add_callout(doc, "Privacy", "The Player QR identifies the campaign but does not contain the Owner PIN, GM PIN, or a permanent Admin credential.")
    doc.add_heading("Publish the current scene", level=2)
    add_numbered(doc, "Enter a short scene title and a player-safe description in Current scene.", num_id)
    add_numbered(doc, "Choose Publish scene. Connected Player screens update immediately.", num_id)
    add_numbered(doc, "Update and republish whenever the shared situation changes. Keep secrets in your own notes; everything in this field is public to the campaign.", num_id)
    doc.add_heading("Before initiative", level=2)
    add_bullet(doc, "Confirm the correct campaign and roster.")
    add_bullet(doc, "Publish the opening scene.")
    add_bullet(doc, "Confirm each Player screen shows Connected and the correct character.")
    add_bullet(doc, "Set table audio and test any RFID function cards.")

    page_break(doc)
    add_section_title(doc, "07", "Run an encounter", "GM live play")
    doc.add_heading("Build the encounter", level=2)
    add_numbered(doc, "Select participating player characters. Enter initiative values or choose Roll initiative.", num_id)
    add_numbered(doc, "Add each NPC or enemy with a name, health value, and initiative.", num_id)
    add_numbered(doc, "Choose Start encounter. Nexus orders the combatants and displays the Initiative board.", num_id)
    doc.add_heading("During play", level=2)
    add_bullet(doc, "Use Next turn and Previous to move through initiative. Round count advances as the order cycles.")
    add_bullet(doc, "Apply damage or healing from the combatant card. Character health is synchronized to the Player view.")
    add_bullet(doc, "Add or remove conditions as effects begin and end.")
    add_bullet(doc, "Edit initiative or reorder combatants when the table needs a correction.")
    add_bullet(doc, "Add a late character or NPC without ending the encounter.")
    add_bullet(doc, "For template-defined trackers, use the GM actions shown when their visibility rule is met. D&D 5e death-save controls appear at zero HP; players see the result but cannot change it.")
    add_callout(doc, "Player signal", "The active character receives a prominent Your turn banner. If a Player screen briefly shows Reconnecting, it falls back to periodic refreshes until the live connection returns.")
    doc.add_heading("Correct or finish", level=2)
    add_bullet(doc, "Reset round returns the encounter to round 1 and the first combatant.")
    add_bullet(doc, "End encounter clears initiative and returns the campaign to exploration mode.")
    add_bullet(doc, "The Owner's emergency session reset is a last resort for a stuck session. It clears live state, not stored characters.")

    page_break(doc)
    add_section_title(doc, "08", "Player directions", "Join and follow the table")
    doc.add_heading("Join your character", level=2)
    add_numbered(doc, "Connect your phone, tablet, or Player Controller to the same network as Nexus.", num_id)
    add_numbered(doc, "Scan the campaign Player QR shown by the GM. If no QR is available, open http://sublim3-nexus.local:3000/player/.", num_id)
    add_numbered(doc, "Choose the campaign, then choose the character assigned to you.", num_id)
    add_numbered(doc, "Choose Join the table. No Owner or GM PIN is required. This device remembers your character.", num_id)
    doc.add_heading("What your screen shows", level=2)
    add_bullet(doc, "Character name, player name, role, level, and defense when supplied by the game system.")
    add_bullet(doc, "Live resources such as health, mana, sanity, stamina, or other campaign-defined values.")
    add_bullet(doc, "Conditions and visible trackers, including synchronized D&D 5e death-save status when applicable.")
    add_bullet(doc, "The current public scene and public notes from the GM.")
    add_bullet(doc, "The active combatant, round number, and a Your turn banner when your character is active.")
    add_callout(doc, "Read-only play", "The Player screen is designed primarily as a live display. Ask the GM to correct health, conditions, notes, or encounter data.")
    doc.add_heading("Switch character or device", level=2)
    doc.add_paragraph("Choose Switch character to clear the local Player session, then select a different campaign and character. On a shared device, switch at the end of play so the next player does not inherit your view.")
    doc.add_heading("If the screen reconnects", level=2)
    doc.add_paragraph("Stay on the Nexus network and leave the Player page open. Nexus automatically polls for updates during a brief interruption and retries the live connection. If the page does not recover, reload it; your saved character selection should return.")

    page_break(doc)
    add_section_title(doc, "09", "Settings and maintenance", "Keep Nexus ready")
    doc.add_heading("Recommended first-session settings", level=2)
    table = doc.add_table(rows=1, cols=2); table.style = "Table Grid"; table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, text in enumerate(("Setting", "Recommended starting point")):
        table.rows[0].cells[i].text = text; set_cell_shading(table.rows[0].cells[i], NAVY)
        for run in table.rows[0].cells[i].paragraphs[0].runs: run.bold = True; run.font.color.rgb = rgb(WHITE)
    rows = [
        ("Maximum volume", "A comfortable ceiling for your room and speakers."),
        ("Startup volume", "Lower than the maximum; 55% is the factory software default."),
        ("Volume change", "5% per function-card scan."),
        ("Stop playout timer", "Off for active sessions; 60–120 minutes for unattended ambience."),
        ("RFID interaction", "Swipe for most tables; Place when removing the card should stop audio."),
        ("Repeat-scan delay", "2 seconds to prevent accidental duplicates."),
    ]
    for label, value in rows:
        cells = table.add_row().cells; cells[0].text = label; cells[1].text = value
        cells[0].paragraphs[0].runs[0].bold = True
    set_table_widths(table, [2700, 6660]); set_repeat_table_header(table.rows[0])
    doc.add_heading("Update Nexus", level=2)
    add_numbered(doc, "Connect Nexus to Home Wi-Fi with internet access.", num_id)
    add_numbered(doc, "Open Settings and choose Update. Confirm the prompt.", num_id)
    add_numbered(doc, "Keep power connected while Nexus downloads, installs, and restarts.", num_id)
    add_numbered(doc, "Wait for the refreshed Settings page and read the success or failure notice.", num_id)
    add_callout(doc, "Do not interrupt", "Closing the browser is usually harmless, but loss of power can damage data or leave the software incomplete.", "warning")
    doc.add_heading("Routine care", level=2)
    add_bullet(doc, "Keep the enclosure dry and dust the exterior with a soft, dry cloth.")
    add_bullet(doc, "Review paired clients and revoke old devices.")
    add_bullet(doc, "Test the network, audio renderer, RFID reader, and Player QR before an important session.")
    add_bullet(doc, "Use Shut Down before unplugging or transporting Nexus.")

    page_break(doc)
    add_section_title(doc, "10", "Troubleshooting", "Fast recovery")
    table = doc.add_table(rows=1, cols=2); table.style = "Table Grid"; table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, text in enumerate(("Symptom", "What to do")):
        table.rows[0].cells[i].text = text; set_cell_shading(table.rows[0].cells[i], NAVY)
        for run in table.rows[0].cells[i].paragraphs[0].runs: run.bold = True; run.font.color.rgb = rgb(WHITE)
    issues = [
        ("SubLim3-Nexus is not listed", "Wait two minutes after power-on. Move closer, toggle Wi-Fi on the client, and scan again. If Nexus was on Home Wi-Fi, look for it on that network."),
        ("Browser says no internet", "This is normal in Local Mode. Keep the Wi-Fi connection and open sublim3-nexus.local:3000."),
        ("The Nexus address does not open", "Confirm the client is on the same network. Try http://sublim3-nexus.local:3000, then the IP address shown by the router or Settings."),
        ("Home Wi-Fi switch failed", "Reconnect to SubLim3-Nexus. Failed Home Wi-Fi connections automatically restore Local Mode."),
        ("Owner access expired", "Enter the Owner recovery PIN again. After five incorrect attempts, wait at least one minute before retrying."),
        ("Guest GM cannot enter", "Confirm the campaign and current GM PIN. The Owner may have rotated the PIN or revoked the session."),
        ("Player cannot find a character", "The Owner must create the character in that campaign. Reload the Player selection after it is saved."),
        ("Player shows Reconnecting", "Keep the page open and verify Wi-Fi. Nexus polls automatically; reload if the connection does not return."),
        ("No audio is heard", "Keep Media open on the renderer device, confirm browser audio permission and speaker output, raise volume, and verify the selected file."),
        ("RFID scan does nothing", "Check the latest reader event, confirm the UID has a saved binding, test the entered UID, and review scan mode and repeat delay."),
        ("Update does not complete", "Keep power connected and wait for Nexus Core to return. Rejoin the network and reopen Settings to read the preserved result notice."),
    ]
    for symptom, fix in issues:
        cells = table.add_row().cells; cells[0].text = symptom; cells[1].text = fix
        cells[0].paragraphs[0].runs[0].bold = True
    set_table_widths(table, [2700, 6660]); set_repeat_table_header(table.rows[0])
    add_callout(doc, "Still stuck?", "Record the Core address, software version, network mode, and the exact on-screen message before contacting support. Never send your recovery PIN or Wi-Fi key.")

    page_break(doc)
    add_section_title(doc, "11", "Quick reference", "At the table")
    doc.add_heading("Local addresses", level=2)
    table = doc.add_table(rows=1, cols=2); table.style = "Table Grid"; table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, text in enumerate(("Screen", "Address")):
        table.rows[0].cells[i].text = text; set_cell_shading(table.rows[0].cells[i], NAVY)
        for run in table.rows[0].cells[i].paragraphs[0].runs: run.bold = True; run.font.color.rgb = rgb(WHITE)
    refs = [
        ("Owner Console", "http://sublim3-nexus.local:3000/"),
        ("GM Console", "http://sublim3-nexus.local:3000/gm/"),
        ("Player", "http://sublim3-nexus.local:3000/player/"),
        ("Media", "http://sublim3-nexus.local:3000/media/"),
        ("RFID Cards", "http://sublim3-nexus.local:3000/rfid/"),
        ("Media Library", "http://sublim3-nexus.local:3000/library/"),
        ("Player Controllers", "http://sublim3-nexus.local:3000/controllers/"),
        ("Settings", "http://sublim3-nexus.local:3000/settings/"),
    ]
    for screen, address in refs:
        cells = table.add_row().cells; cells[0].text = screen; cells[1].text = address
        cells[0].paragraphs[0].runs[0].bold = True
    set_table_widths(table, [2700, 6660]); set_repeat_table_header(table.rows[0])
    doc.add_heading("Session close-out", level=2)
    add_bullet(doc, "End the live encounter and publish a closing scene if desired.")
    add_bullet(doc, "Stop audio and remove RFID cards from a Place-mode reader.")
    add_bullet(doc, "Unpair temporary guest GM devices and shared Player devices when appropriate.")
    add_bullet(doc, "Open Settings and choose Shut Down. Wait for the Core to stop before disconnecting power.")
    doc.add_heading("Security-card reminder", level=2)
    add_callout(doc, "Keep private", "Owner recovery PIN and Local Wi-Fi key. Share the guest GM PIN only with a trusted GM, and prefer the campaign Player QR for players.", "caution")
    doc.add_paragraph("SubLim3 Nexus is designed to keep play moving locally, even when the internet does not. Keep this manual and the security card in a safe place near the unit.")
    closing = doc.add_paragraph("READY WHEN THE TABLE IS")
    closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
    closing.paragraph_format.space_before = Pt(28)
    closing.runs[0].font.name = "Aptos Display"; closing.runs[0].font.size = Pt(15); closing.runs[0].bold = True; closing.runs[0].font.color.rgb = rgb(CYAN)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
